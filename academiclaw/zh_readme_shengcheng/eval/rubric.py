"""
pengyichen-query2 评分标准 (Rubric)
任务：Web项目README文档生成 - MYmajorCS-Web

总分：100分

评分维度：
一、文件交付与基础结构 (15分)
  1. README.md 文件存在且非空 (3分)
  2. 七个必需章节完整性 (7分): 项目概述/技术栈/项目结构/安装部署/使用说明/接口文档/用例文档 各1分
  3. Markdown 格式规范性 (5分): 代码块(2分) + 表格(2分) + 层次结构(1分)

二、文档内容质量 (25分) - LLM 评估
  1. 技术栈描述准确性 (7分) - 与 package.json / 项目结构一致
  2. 项目结构描述准确性 (7分) - 与 root_directory.txt 一致
  3. 部署流程可行性 (6分) - 步骤清晰可执行, Docker 说明与 Dockerfile 一致
  4. 功能描述完整性 (5分) - 覆盖主要功能模块

三、接口文档质量 (30分) - LLM 评估 (对比 接口设计.docx)
  1. 接口覆盖度 (15分)
  2. 参数描述准确性 (7分)
  3. 示例代码完整性 (5分)
  4. 错误处理描述 (3分)

四、用例文档质量 (30分) - LLM 评估 (对比 用例设计.docx)
  1. 用例覆盖度 (15分)
  2. 用例结构完整性 (7分)
  3. 异常流程覆盖 (3分)
  4. 业务逻辑准确性 (5分)
"""

import os
import re
import json
from typing import Tuple, Dict, Any, List

try:
    import openai
except ImportError:
    openai = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


# =============================================================================
# 环境配置与 LLM 工具
# =============================================================================

def _load_env(answer_dir: str) -> dict:
    """从 answer_dir 和 query 根目录加载 .env 配置"""
    values = {}
    for env_dir in [answer_dir, os.path.join(os.path.dirname(__file__), "..")]:
        env_path = os.path.join(env_dir, ".env")
        if os.path.exists(env_path):
            with open(env_path, "r") as f:
                for line in f:
                    line = line.strip()
                    if not line or line.startswith("#") or "=" not in line:
                        continue
                    k, v = line.split("=", 1)
                    if k.strip() not in values:
                        values[k.strip()] = v.strip().strip("'\"")
    return values


def _get_text_eval_config(answer_dir: str) -> dict:
    """获取文本评估 LLM 配置"""
    env = _load_env(answer_dir)

    def g(key, default=""):
        return os.environ.get(key) or env.get(key) or default

    return {
        "api_key": g("EVAL_TEXT_API_KEY", g("ANTHROPIC_API_KEY")),
        "api_base": g("EVAL_TEXT_API_BASE_URL", g("ANTHROPIC_BASE_URL")),
        "model": g("EVAL_TEXT_MODEL", "openai/gpt-5.2"),
    }


def _call_llm_judge(prompt: str, config: dict) -> str:
    """调用 LLM 进行文本评估"""
    if not openai or not config.get("api_key"):
        return ""
    try:
        base = config["api_base"].rstrip("/")
        if not base.endswith("/v1"):
            base += "/v1"
        client = openai.OpenAI(api_key=config["api_key"], base_url=base)
        resp = client.chat.completions.create(
            model=config["model"],
            messages=[{"role": "user", "content": prompt}],
            max_tokens=4096,
            temperature=0,
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        print("[RUBRIC] LLM Judge 调用失败: %s" % e)
        return ""


def _parse_llm_json(text: str) -> dict:
    """从 LLM 响应中解析 JSON"""
    if not text:
        return {}
    # 尝试从 markdown code block 中提取
    m = re.search(r"```json\s*(.*?)\s*```", text, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(1))
        except json.JSONDecodeError:
            pass
    # 尝试直接提取 JSON object
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        try:
            return json.loads(text[start : end + 1])
        except json.JSONDecodeError:
            pass
    return {}


# =============================================================================
# 辅助函数
# =============================================================================

def _read_file(path: str) -> str:
    """安全读取文本文件"""
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception:
        return ""


def _find_readme(answer_dir: str) -> str:
    """在 answer_dir 中查找 README 文件（优先根目录，再递归子目录）"""
    # 优先查找根目录
    for name in ("README.md", "readme.md", "Readme.md", "README.MD"):
        p = os.path.join(answer_dir, name)
        if os.path.isfile(p):
            return p
    # 递归查找（排除隐藏目录和常见无关目录）
    for root, dirs, files in os.walk(answer_dir):
        dirs[:] = [d for d in dirs if not d.startswith(".")
                   and d not in ("__pycache__", "node_modules", "context", ".sii")]
        for f in files:
            if f.lower() == "readme.md":
                return os.path.join(root, f)
    return ""


def _load_docx_text(docx_path: str) -> str:
    """从 .docx 文件提取纯文本（含表格）"""
    if DocxDocument is None:
        return ""
    try:
        doc = DocxDocument(docx_path)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception:
        return ""


def _find_context_file(answer_dir: str, filename: str) -> str:
    """在 context 目录中查找参考文件"""
    # 先查 query 根目录的 context/
    query_root = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..")
    for base in [answer_dir, query_root]:
        p = os.path.join(base, "context", filename)
        if os.path.exists(p):
            return p
    return ""


def _extract_section(content: str, section_keyword: str) -> str:
    """从 Markdown 中提取包含指定关键词的章节内容"""
    lines = content.split("\n")
    start_idx = -1
    start_level = 0

    for i, line in enumerate(lines):
        m = re.match(r"^(#{1,3})\s+", line)
        if m and section_keyword in line:
            start_idx = i
            start_level = len(m.group(1))
            break

    if start_idx == -1:
        # 降级：直接搜索关键词位置
        idx = content.find(section_keyword)
        if idx == -1:
            return ""
        rest = content[idx:]
        next_h = re.search(r"\n#{1,2}\s", rest[len(section_keyword):])
        if next_h:
            return rest[: len(section_keyword) + next_h.start()]
        return rest

    # 找到下一个同级或更高级标题
    section_lines = [lines[start_idx]]
    for i in range(start_idx + 1, len(lines)):
        m = re.match(r"^(#{1,3})\s+", lines[i])
        if m and len(m.group(1)) <= start_level:
            break
        section_lines.append(lines[i])

    return "\n".join(section_lines)


# =============================================================================
# 一、文件交付与基础结构 (15分)
# =============================================================================

def _check_structure(answer_dir: str) -> Tuple[int, Dict[str, Any]]:
    score = 0
    details = {}
    issues = []

    readme_path = _find_readme(answer_dir)

    # 1.1 README.md 存在且非空 (3分)
    if not readme_path:
        details["1.1 文件存在 (3)"] = "0/3 - 未找到 README.md"
        issues.append("缺少 README.md")
        return 0, {"score": 0, "details": details, "issues": issues}

    content = _read_file(readme_path)
    if len(content.strip()) < 100:
        score += 1
        details["1.1 文件存在 (3)"] = "1/3 - 文件存在但内容过少 (%d 字符)" % len(content)
        issues.append("README 内容不足 100 字符")
        return score, {"score": score, "details": details, "issues": issues}

    score += 3
    details["1.1 文件存在 (3)"] = "3/3 - 文件存在且内容充实 (%d 字符)" % len(content)

    # 1.2 七个必需章节 (7分，每个 1 分)
    required_sections = [
        "项目概述", "技术栈", "项目结构", "安装部署", "使用说明", "接口文档", "用例文档",
    ]
    # 宽松匹配别名
    aliases = {
        "项目概述": ["项目概述", "项目简介", "概述", "overview"],
        "技术栈": ["技术栈", "技术架构", "tech stack", "technology"],
        "项目结构": ["项目结构", "目录结构", "project structure"],
        "安装部署": ["安装部署", "安装与部署", "安装", "部署", "installation", "deploy"],
        "使用说明": ["使用说明", "使用指南", "用法", "usage"],
        "接口文档": ["接口文档", "接口说明", "api", "API"],
        "用例文档": ["用例文档", "用例", "use case"],
    }
    section_score = 0
    section_detail = {}
    content_lower = content.lower()
    for sec in required_sections:
        found = False
        for alias in aliases.get(sec, [sec]):
            if alias.lower() in content_lower:
                found = True
                break
        if found:
            section_score += 1
            section_detail[sec] = "1/1"
        else:
            section_detail[sec] = "0/1 - 缺失"
            issues.append("缺少「%s」章节" % sec)
    score += section_score
    details["1.2 章节完整性 (7)"] = section_detail

    # 1.3 格式规范性 (5分)
    fmt_score = 0
    fmt_detail = {}

    # 代码块 (2分): 有语法高亮的代码块 2 分; 有代码块但无高亮 1 分
    if re.search(r"```(?:bash|json|http|javascript|typescript|shell|sh|sql|html|css|yaml|dockerfile)", content, re.IGNORECASE):
        fmt_score += 2
        fmt_detail["代码块 (2)"] = "2/2 - 含语法高亮代码块"
    elif "```" in content:
        fmt_score += 1
        fmt_detail["代码块 (2)"] = "1/2 - 有代码块但缺语法标注"
    else:
        fmt_detail["代码块 (2)"] = "0/2 - 无代码块"
        issues.append("缺少代码块")

    # 表格 (2分)
    if "|" in content and "---" in content:
        fmt_score += 2
        fmt_detail["表格 (2)"] = "2/2 - 含 Markdown 表格"
    else:
        fmt_detail["表格 (2)"] = "0/2 - 无表格"
        issues.append("缺少 Markdown 表格")

    # 层次结构 (1分): 同时有 ## 和 ###
    if "## " in content and "### " in content:
        fmt_score += 1
        fmt_detail["层次结构 (1)"] = "1/1"
    elif "## " in content:
        fmt_detail["层次结构 (1)"] = "0/1 - 仅有二级标题"
    else:
        fmt_detail["层次结构 (1)"] = "0/1 - 标题层次不清"
        issues.append("标题层次不清")

    score += fmt_score
    details["1.3 格式规范 (5)"] = fmt_detail

    return min(score, 15), {"score": min(score, 15), "details": details, "issues": issues}


# =============================================================================
# 二、文档内容质量 (25分) - LLM 评估
# =============================================================================

def _check_content_quality(answer_dir: str) -> Tuple[int, Dict[str, Any]]:
    readme_path = _find_readme(answer_dir)
    if not readme_path:
        return 0, {"score": 0, "details": {"error": "未找到 README.md"}, "issues": []}

    content = _read_file(readme_path)
    if not content:
        return 0, {"score": 0, "details": {"error": "README 为空"}, "issues": []}

    # 组装参考信息
    ref_info = ""
    root_dir_path = _find_context_file(answer_dir, "root_directory.txt")
    if root_dir_path:
        t = _read_file(root_dir_path)
        if t:
            ref_info += "== 项目目录结构 ==\n%s\n\n" % t

    pkg_path = _find_context_file(answer_dir, "package.json")
    if pkg_path:
        t = _read_file(pkg_path)
        if t:
            ref_info += "== package.json ==\n%s\n\n" % t

    docker_path = _find_context_file(answer_dir, "dockerfilie")
    if docker_path:
        t = _read_file(docker_path)
        if t:
            ref_info += "== Dockerfile (前 2000 字符) ==\n%s\n" % t[:2000]

    config = _get_text_eval_config(answer_dir)
    prompt = """你是一位严格的技术文档评审专家。请评估以下 Web 项目 README 文档的内容质量。

## 项目真实信息（用于校验文档准确性）
%s

## 待评估的 README 文档（前 8000 字符）
%s

## 评分维度（请严格按分数区间打分）

### 技术栈描述准确性 (0-7分)
- 前端：是否准确提到 Vue 3、@heroicons/vue
- 后端：Node.js 相关技术是否与目录结构(controllers/models/routes/middleware/utils)一致
- 基础设施：Docker/NGINX 是否正确描述
- 7=完全准确且详细, 4-6=基本准确有小瑕疵, 1-3=有明显错误, 0=完全缺失

### 项目结构描述准确性 (0-7分)
- 目录树是否与 root_directory.txt 一致（backend, webApp, AspireHost 三个主要模块）
- 各目录功能描述是否正确
- 7=完全匹配且描述清楚, 4-6=大部分匹配, 1-3=有重大偏差, 0=完全缺失

### 部署流程可行性 (0-6分)
- 是否包含环境要求、依赖安装、启动命令
- Docker 部署说明是否与 Dockerfile 一致
- 6=步骤完整清晰可执行, 3-5=基本可用, 1-2=不够具体, 0=缺失

### 功能描述完整性 (0-5分)
- 使用说明是否涵盖主要功能（课程浏览、搜索、学习计划、账户等）
- 5=全面详细, 3-4=基本覆盖, 1-2=粗略, 0=缺失

请严格按以下 JSON 格式输出（不要包含任何其他内容）：
```json
{
    "tech_stack": {"score": 0, "reason": ""},
    "project_structure": {"score": 0, "reason": ""},
    "deployment": {"score": 0, "reason": ""},
    "features": {"score": 0, "reason": ""}
}
```""" % (ref_info, content[:8000])

    llm_text = _call_llm_judge(prompt, config)
    result = _parse_llm_json(llm_text)

    if result:
        return _score_content_llm(result)
    return _score_content_fallback(content)


def _score_content_llm(result: dict) -> Tuple[int, Dict[str, Any]]:
    score = 0
    details = {}
    mapping = [
        ("tech_stack", "技术栈准确性", 7),
        ("project_structure", "项目结构准确性", 7),
        ("deployment", "部署流程可行性", 6),
        ("features", "功能描述完整性", 5),
    ]
    for key, label, max_s in mapping:
        d = result.get(key, {})
        raw = d.get("score", 0)
        if isinstance(raw, str):
            try:
                raw = int(raw)
            except ValueError:
                raw = 0
        s = max(0, min(int(raw), max_s))
        score += s
        details["%s (%d)" % (label, max_s)] = "%d/%d - %s" % (s, max_s, d.get("reason", ""))
    score = min(score, 25)
    return score, {"score": score, "details": details, "issues": []}


def _score_content_fallback(content: str) -> Tuple[int, Dict[str, Any]]:
    """LLM 不可用时的基础关键词评估"""
    score = 0
    details = {}

    # 技术栈关键词
    tech_kws = ["vue", "vue 3", "@heroicons/vue", "node.js", "express",
                "docker", "nginx", "jwt", "vite"]
    tech_found = sum(1 for kw in tech_kws if kw.lower() in content.lower())
    tech_s = min(tech_found, 7)
    score += tech_s
    details["技术栈关键词"] = "%d/7 - 匹配 %d 个" % (tech_s, tech_found)

    # 项目结构关键词
    struct_kws = {"backend": 2, "webApp": 2, "AspireHost": 1, "controllers": 1, "models": 1}
    struct_s = min(sum(v for k, v in struct_kws.items() if k in content), 7)
    score += struct_s
    details["项目结构关键词"] = "%d/7" % struct_s

    # 部署关键词
    deploy_kws = ["npm install", "docker", "git clone", "npm run"]
    deploy_found = sum(1 for kw in deploy_kws if kw.lower() in content.lower())
    deploy_s = min(deploy_found * 2, 6)
    score += deploy_s
    details["部署关键词"] = "%d/6" % deploy_s

    # 长度兜底
    if len(content) >= 3000:
        score += 3
    elif len(content) >= 1500:
        score += 1
    details["文档长度"] = "%d 字符" % len(content)

    score = min(score, 25)
    details["注意"] = "LLM 不可用，使用基础关键词评估，分数可能不准确"
    return score, {"score": score, "details": details, "issues": []}


# =============================================================================
# 三、接口文档质量 (30分) - LLM 评估
# =============================================================================

def _check_interface_docs(answer_dir: str) -> Tuple[int, Dict[str, Any]]:
    readme_path = _find_readme(answer_dir)
    if not readme_path:
        return 0, {"score": 0, "details": {"error": "未找到 README"}, "issues": []}

    content = _read_file(readme_path)
    if not content:
        return 0, {"score": 0, "details": {"error": "README 为空"}, "issues": []}

    section = _extract_section(content, "接口文档")
    if not section:
        section = _extract_section(content, "接口")
    if not section:
        section = _extract_section(content, "API")
    if not section or len(section.strip()) < 50:
        return 0, {"score": 0, "details": {"error": "接口文档章节缺失或内容过少"}, "issues": ["接口文档缺失"]}

    # 加载参考文档
    ref_text = ""
    ref_path = _find_context_file(answer_dir, "接口设计.docx")
    if ref_path:
        ref_text = _load_docx_text(ref_path)

    config = _get_text_eval_config(answer_dir)
    prompt = """你是一位严格的 API 文档评审专家。请评估以下接口文档的质量。

## 参考接口设计文档（真实接口信息，用于对比准确性）
%s

## 待评估的接口文档章节（前 6000 字符）
%s

## 评分维度

### 接口覆盖度 (0-15分)
- 是否覆盖了主要接口类别（认证/登录注册、用户管理、课程/专业查询、学习计划、组织/战队等）
- 每个类别下是否有具体的接口路径和方法(GET/POST/PUT/DELETE)
- 15=全面覆盖参考文档中的主要接口, 10-14=覆盖大部分, 5-9=仅覆盖少数, 0-4=几乎没有

### 参数描述准确性 (0-7分)
- 请求参数和响应参数是否有清晰描述
- 参数类型、是否必填等信息是否完整
- 与参考文档一致性
- 7=准确完整, 4-6=基本准确, 1-3=不完整, 0=缺失

### 示例代码完整性 (0-5分)
- 是否包含请求示例（curl 或代码）
- 是否包含响应 JSON 示例
- 5=丰富完整, 3-4=有部分示例, 1-2=很少, 0=无

### 错误处理描述 (0-3分)
- 是否说明了常见错误码/错误响应格式
- 3=详细, 1-2=有提及, 0=无

请严格按以下 JSON 格式输出：
```json
{
    "coverage": {"score": 0, "reason": ""},
    "params": {"score": 0, "reason": ""},
    "examples": {"score": 0, "reason": ""},
    "errors": {"score": 0, "reason": ""}
}
```""" % (
        ref_text[:4000] if ref_text else "（参考文档不可用，请根据 Web 项目常见接口标准评估）",
        section[:6000],
    )

    llm_text = _call_llm_judge(prompt, config)
    result = _parse_llm_json(llm_text)

    if result:
        return _score_interface_llm(result)
    return _score_interface_fallback(section)


def _score_interface_llm(result: dict) -> Tuple[int, Dict[str, Any]]:
    score = 0
    details = {}
    mapping = [
        ("coverage", "接口覆盖度", 15),
        ("params", "参数描述准确性", 7),
        ("examples", "示例代码完整性", 5),
        ("errors", "错误处理描述", 3),
    ]
    for key, label, max_s in mapping:
        d = result.get(key, {})
        raw = d.get("score", 0)
        if isinstance(raw, str):
            try:
                raw = int(raw)
            except ValueError:
                raw = 0
        s = max(0, min(int(raw), max_s))
        score += s
        details["%s (%d)" % (label, max_s)] = "%d/%d - %s" % (s, max_s, d.get("reason", ""))
    score = min(score, 30)
    return score, {"score": score, "details": details, "issues": []}


def _score_interface_fallback(section: str) -> Tuple[int, Dict[str, Any]]:
    """LLM 不可用时的关键词评估"""
    score = 0
    details = {}
    sec_lower = section.lower()

    # 覆盖度：接口类别关键词
    api_kws = ["登录", "注册", "认证", "auth", "用户", "user", "课程", "course",
               "计划", "plan", "搜索", "search", "通知", "战队", "team"]
    found = sum(1 for kw in api_kws if kw.lower() in sec_lower)
    cov_s = min(found * 2, 15)
    score += cov_s
    details["接口覆盖度"] = "%d/15 - 匹配 %d 个类别" % (cov_s, found)

    # 参数描述
    param_kws = ["请求参数", "响应", "request", "response", "string", "number", "必填", "required"]
    param_found = sum(1 for kw in param_kws if kw.lower() in sec_lower)
    param_s = min(param_found, 7)
    score += param_s
    details["参数描述"] = "%d/7" % param_s

    # 示例
    ex_s = 0
    if "```json" in section:
        ex_s += 3
    if re.search(r"(POST|GET|PUT|DELETE)\s+/api/", section):
        ex_s += 2
    ex_s = min(ex_s, 5)
    score += ex_s
    details["示例代码"] = "%d/5" % ex_s

    # 错误处理
    err_kws = ["错误", "error", "失败", "状态码", "status"]
    err_s = min(sum(1 for kw in err_kws if kw.lower() in sec_lower), 3)
    score += err_s
    details["错误处理"] = "%d/3" % err_s

    score = min(score, 30)
    details["注意"] = "LLM 不可用，使用基础关键词评估"
    return score, {"score": score, "details": details, "issues": []}


# =============================================================================
# 四、用例文档质量 (30分) - LLM 评估
# =============================================================================

def _check_usecase_docs(answer_dir: str) -> Tuple[int, Dict[str, Any]]:
    readme_path = _find_readme(answer_dir)
    if not readme_path:
        return 0, {"score": 0, "details": {"error": "未找到 README"}, "issues": []}

    content = _read_file(readme_path)
    if not content:
        return 0, {"score": 0, "details": {"error": "README 为空"}, "issues": []}

    section = _extract_section(content, "用例文档")
    if not section:
        section = _extract_section(content, "用例")
    if not section or len(section.strip()) < 50:
        return 0, {"score": 0, "details": {"error": "用例文档章节缺失或内容过少"}, "issues": ["用例文档缺失"]}

    # 加载参考文档
    ref_text = ""
    ref_path = _find_context_file(answer_dir, "用例设计.docx")
    if ref_path:
        ref_text = _load_docx_text(ref_path)

    config = _get_text_eval_config(answer_dir)
    prompt = """你是一位严格的软件测试和需求分析专家。请评估以下用例文档的质量。

## 参考用例设计文档（真实用例信息，用于对比准确性）
%s

## 待评估的用例文档章节（前 6000 字符）
%s

## 评分维度

### 用例覆盖度 (0-15分)
- 是否覆盖了主要功能场景（用户认证、课程浏览、搜索、学习计划管理、管理员操作等）
- 用例数量是否充足（3-5 个为基本，>5 为充分）
- 15=全面覆盖参考文档中的主要场景, 10-14=覆盖大部分, 5-9=仅覆盖少数, 0-4=几乎没有

### 用例结构完整性 (0-7分)
- 每个用例是否包含：参与者、前置条件、主流程步骤、预期结果
- 流程步骤是否具体清晰（而非含糊笼统）
- 7=结构完整规范, 4-6=基本完整, 1-3=结构不全, 0=无结构

### 异常流程覆盖 (0-3分)
- 是否包含备选流/异常流（如：登录失败、数据校验不通过、权限不足等）
- 3=多个用例有异常流, 1-2=少数有, 0=无

### 业务逻辑准确性 (0-5分)
- 用例描述的业务逻辑是否符合一个培养方案/选课规划 Web 项目的实际场景
- 与参考文档的一致性
- 5=准确合理, 3-4=基本合理, 1-2=有偏差, 0=不相关

请严格按以下 JSON 格式输出：
```json
{
    "coverage": {"score": 0, "reason": ""},
    "structure": {"score": 0, "reason": ""},
    "exceptions": {"score": 0, "reason": ""},
    "accuracy": {"score": 0, "reason": ""}
}
```""" % (
        ref_text[:4000] if ref_text else "（参考文档不可用，请根据 Web 项目常见用例标准评估）",
        section[:6000],
    )

    llm_text = _call_llm_judge(prompt, config)
    result = _parse_llm_json(llm_text)

    if result:
        return _score_usecase_llm(result)
    return _score_usecase_fallback(section)


def _score_usecase_llm(result: dict) -> Tuple[int, Dict[str, Any]]:
    score = 0
    details = {}
    mapping = [
        ("coverage", "用例覆盖度", 15),
        ("structure", "用例结构完整性", 7),
        ("exceptions", "异常流程覆盖", 3),
        ("accuracy", "业务逻辑准确性", 5),
    ]
    for key, label, max_s in mapping:
        d = result.get(key, {})
        raw = d.get("score", 0)
        if isinstance(raw, str):
            try:
                raw = int(raw)
            except ValueError:
                raw = 0
        s = max(0, min(int(raw), max_s))
        score += s
        details["%s (%d)" % (label, max_s)] = "%d/%d - %s" % (s, max_s, d.get("reason", ""))
    score = min(score, 30)
    return score, {"score": score, "details": details, "issues": []}


def _score_usecase_fallback(section: str) -> Tuple[int, Dict[str, Any]]:
    """LLM 不可用时的关键词评估"""
    score = 0
    details = {}

    # 覆盖度
    uc_kws = ["注册", "登录", "认证", "课程", "搜索", "计划", "管理",
              "浏览", "创建", "查看", "编辑", "删除"]
    found = sum(1 for kw in uc_kws if kw in section)
    cov_s = min(found * 2, 15)
    score += cov_s
    details["用例覆盖度"] = "%d/15 - 匹配 %d 个场景" % (cov_s, found)

    # 结构完整性
    struct_kws = ["参与者", "前置条件", "主流程", "步骤", "预期结果", "后置条件", "主成功场景"]
    struct_found = sum(1 for kw in struct_kws if kw in section)
    struct_s = min(struct_found * 2, 7)
    score += struct_s
    details["结构完整性"] = "%d/7" % struct_s

    # 异常流程
    exc_kws = ["异常", "备选", "失败", "错误", "替代流", "边界"]
    exc_s = min(sum(1 for kw in exc_kws if kw in section), 3)
    score += exc_s
    details["异常流程"] = "%d/3" % exc_s

    # 用例数量
    uc_count = len(re.findall(r"用例\s*(?:UC[-_]?)?\d+|UC[-_]?\d+", section, re.IGNORECASE))
    if uc_count >= 5:
        biz_s = 5
    elif uc_count >= 3:
        biz_s = 3
    elif uc_count >= 1:
        biz_s = 1
    else:
        biz_s = 0
    score += biz_s
    details["业务逻辑(用例数量)"] = "%d/5 - 发现 %d 个用例" % (biz_s, uc_count)

    score = min(score, 30)
    details["注意"] = "LLM 不可用，使用基础关键词评估"
    return score, {"score": score, "details": details, "issues": []}


# =============================================================================
# 主评估函数
# =============================================================================

def evaluate(answer_dir: str) -> Tuple[int, Dict[str, Any]]:
    """
    评估 agent 的输出。

    Args:
        answer_dir: agent 输出目录的绝对路径

    Returns:
        (score, report) — score 0-100, report 包含详细评估报告
    """
    s1, r1 = _check_structure(answer_dir)
    s2, r2 = _check_content_quality(answer_dir)
    s3, r3 = _check_interface_docs(answer_dir)
    s4, r4 = _check_usecase_docs(answer_dir)

    total = min(s1 + s2 + s3 + s4, 100)

    report = {
        "total": total,
        "breakdown": {
            "文件交付与基础结构": "%d/15" % s1,
            "文档内容质量": "%d/25" % s2,
            "接口文档质量": "%d/30" % s3,
            "用例文档质量": "%d/30" % s4,
        },
        "dim1_structure": r1,
        "dim2_content": r2,
        "dim3_interface": r3,
        "dim4_usecase": r4,
        "comment": "",
    }

    if total >= 90:
        report["comment"] = "优秀！README 文档内容完整、格式规范、接口和用例文档质量高。"
    elif total >= 75:
        report["comment"] = "良好。文档整体质量不错，个别方面可进一步完善。"
    elif total >= 60:
        report["comment"] = "及格。完成了基本要求，但内容深度或覆盖面有待提升。"
    elif total >= 40:
        report["comment"] = "部分完成。文档有一定内容但存在明显不足。"
    else:
        report["comment"] = "不及格。文档缺失严重或内容质量较差。"

    return total, report


def print_report(score: int, report: Dict[str, Any]) -> None:
    """打印格式化的评分报告"""
    print("=" * 70)
    print("pengyichen-query2 评分报告")
    print("任务：Web项目README文档生成 - MYmajorCS-Web")
    print("=" * 70)
    print("\n总分：%d/100\n" % score)

    breakdown = report.get("breakdown", {})
    if breakdown:
        print("分项得分:")
        for k, v in breakdown.items():
            print("  %s: %s" % (k, v))

    dim_keys = [
        ("dim1_structure", "一、文件交付与基础结构 (15分)"),
        ("dim2_content", "二、文档内容质量 (25分)"),
        ("dim3_interface", "三、接口文档质量 (30分)"),
        ("dim4_usecase", "四、用例文档质量 (30分)"),
    ]

    for key, title in dim_keys:
        dim = report.get(key, {})
        dim_score = dim.get("score", 0)
        print("\n" + "-" * 50)
        print("【%s】 得分: %s" % (title, dim_score))
        print("-" * 50)

        for item_name, item_val in dim.get("details", {}).items():
            if isinstance(item_val, dict):
                print("  %s:" % item_name)
                for k, v in item_val.items():
                    v_str = str(v)
                    if len(v_str) > 120:
                        v_str = v_str[:120] + "..."
                    print("    %s: %s" % (k, v_str))
            else:
                v_str = str(item_val)
                if len(v_str) > 120:
                    v_str = v_str[:120] + "..."
                print("  %s: %s" % (item_name, v_str))

        issues = dim.get("issues", [])
        if issues:
            print("  扣分项:")
            for i, issue in enumerate(issues, 1):
                print("    %d. %s" % (i, issue))

    print("\n" + "=" * 50)
    print("评语：%s" % report.get("comment", ""))
    print("=" * 70)


if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        test_dir = sys.argv[1]
    else:
        test_dir = os.path.join(os.path.dirname(__file__), "..", "gpt-5", "attempt_1")

    if not os.path.isabs(test_dir):
        test_dir = os.path.join(os.getcwd(), test_dir)

    if os.path.exists(test_dir):
        print("正在评估目录: %s\n" % test_dir)
        s, r = evaluate(test_dir)
        print_report(s, r)
    else:
        print("目录不存在: %s" % test_dir)
    sys.exit(0)
